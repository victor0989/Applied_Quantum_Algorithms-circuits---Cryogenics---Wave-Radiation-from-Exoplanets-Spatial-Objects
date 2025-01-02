from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Scaling Laws and Quantum Algorithm', 0)

# Adding Scaling Laws Section
doc.add_heading('Scaling Laws and Emission Ratio Computation', level=1)

doc.add_paragraph('The kinetic and magnetic energy fluxes from Jupiter\'s magnetosphere intercepted by Io and Ganymede '
                  'can be computed from the plasma parameters in the satellites\' vicinity, as measured by the Voyager '
                  'and Galileo spacecraft (Kivelson et al., 2004; Zarka, 2007):')

doc.add_paragraph('P_kin = (V^2) / (R_obs^2)  (3)')
doc.add_paragraph('P_mag = (B^2) / (R_obs^2)  (4)')

doc.add_paragraph('where B is the (sub-)corotating Jovian plasma density and magnetic field amplitude at the satellite '
                  'orbit, and V is the flow velocity:')

doc.add_paragraph('V = V_corot * V_orb = sqrt(G * M_J / L_RJ)  (5)')

doc.add_paragraph('where J, M_J, and R_J characterize Jupiter\'s rotation, mass, and radius, respectively, and L_RJ is '
                  'the satellite orbital radius. R_obs^2 is the cross-section of the obstacle.')

# Adding Quantum Algorithm Section
doc.add_heading('Quantum Algorithm for Optimization', level=1)

doc.add_paragraph('1. (x; y; τ; θ; s; κ) ← (e; 0; 1; 1; e; 1) /* Initialize on central path */')
doc.add_paragraph('2. μ ← 1, σ ← 1 − 1 / (20√2), γ ← 1/10 /* Set parameters */')
doc.add_paragraph('3. while μ ≥ : /* Follow central path until duality gap less than */')
doc.add_paragraph('4. G ← ( 0   A^T   −c    ¯c   I   0   \n'
                  '    −A    0    b    ¯b   0   0  \n'
                  '    c^T  −b^T   0    −¯z   0   1  \n'
                  '   −¯c^T  ¯b^T   ¯z    0   0   0  \n'
                  '    S    0    0    0   X   0  \n'
                  '    0    0    κ    0   0   τ ) /* from eqs.(19) and (22) */')
doc.add_paragraph('5. h ← ( −A^T y + c τ − ¯c θ − s   \n'
                  '       Ax − b τ + ¯b θ    \n'
                  '       −c^T x + b^T y + ¯z θ \n'
                  '       ¯c^T x − ¯b^T y − ¯z τ  \n'
                  '       σ μ e − ˜X˜S e \n'
                  '       σ μ − κ τ ) /* Matrix-vector multiplication performed classically */')
doc.add_paragraph('6. for j = 1,..., L: /* Preconditioning via row normalization */')
doc.add_paragraph('7. g ← k / |G_jk|^2 /* Norm of jth row of G */')
doc.add_paragraph('8. h_j ← h_j / g')
doc.add_paragraph('9. for k = 1,..., L:')
doc.add_paragraph('10. G_jk ← G_jk / g')
doc.add_paragraph('11. Classically compute L2 angles and gate decompositions necessary to perform block-encoding of G and state-preparation of |h (see Ref. [33])')
doc.add_paragraph('12. ξ ← 1')
doc.add_paragraph('13. repeat /* Try smaller and smaller ξ until central path is found */')
doc.add_paragraph('14. ξ ← ξ / 2')
doc.add_paragraph('15. (Δx; Δy; Δτ; Δθ; Δs; Δκ) ← ApprSolve(G, h, ξ)')
doc.add_paragraph('16. (steplength) ← μ(σ − 1)(r + 1) ((Δx)^T s + (Δs)^T x + (Δκ) τ + (Δτ) κ)')
doc.add_paragraph('17. (x; y; τ; θ; s; κ) ← (x; y; τ; θ; s; κ) + (steplength)·(Δx; Δy; Δτ; Δθ; Δs; Δκ)')
doc.add_paragraph('18. until (x; y; τ; θ; s; κ) ∈ N(γ)')
doc.add_paragraph('19. (x; y; τ; θ; s; κ) ← (x; y; τ; θ; s; κ)')
doc.add_paragraph('20. μ ← σ μ')
doc.add_paragraph('21. return x / τ')

# Adding Approximation Solver Section
doc.add_heading('Approximation Solver (ApprSolve)', level=1)

doc.add_paragraph('22. def ApprSolve(G, h, ξ):')
doc.add_paragraph('23. L ← 2N + K + 3')
doc.add_paragraph('24. δ ← 0.1')
doc.add_paragraph('25. ε ← 0.9 ξ')
doc.add_paragraph('26. k ← 57.5 L ln(6L / δ) / (ε^2(1 − ε^2 / 4))')
doc.add_paragraph('27. Run tomography as described in section IV D using k applications and k controlled applications of the QLSS algorithm on the system (G, h)')
doc.add_paragraph('28. return Vector v for which v = 1 and v − v ≤ ξ with probability at least (1 − δ), where v ∝ G^−1 h')

# Save the document
doc_path = '/mnt/data/scaling_laws_and_quantum_algorithm.docx'
doc.save(doc_path)

doc_path
